"""Output transformation tool for converting cached results to different formats."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from pydantic import BaseModel, Field

from ..session_state import SessionState


class OutputTransformerInput(BaseModel):
    """Schema for the output transformer tool."""

    source_type: str = Field(
        description="Type of cached output to transform: 'table' or 'chart'",
        default="table",
    )
    target_format: str = Field(
        description="Target format: 'csv', 'excel', 'pdf', 'docx', 'png', 'json'",
    )
    description: str = Field(
        description="Brief description of the transformation",
        default="Converting cached output",
    )


class OutputTransformerTool:
    """Tool for transforming cached outputs without re-executing code."""

    name = "transform_output"
    description = (
        "Transform previously generated outputs (tables or charts) to different formats "
        "WITHOUT re-running the analysis or re-parsing CSV files. "
        "Use this when user asks to convert, export, or save the last result in a different format. "
        "For example: 'convert that to PDF', 'save the last table as Excel', 'export as DOCX'."
    )

    @staticmethod
    def get_schema() -> Dict[str, Any]:
        """Return OpenAI function calling schema."""
        return {
            "type": "function",
            "function": {
                "name": OutputTransformerTool.name,
                "description": OutputTransformerTool.description,
                "parameters": OutputTransformerInput.model_json_schema(),
            },
        }

    @staticmethod
    def execute(
        arguments: Dict[str, Any], session: SessionState
    ) -> Dict[str, Any]:
        """Execute the transformation."""
        try:
            input_data = OutputTransformerInput(**arguments)
        except Exception as e:
            return {
                "success": False,
                "error": f"Invalid arguments: {e}",
            }

        # Get cached result
        if input_data.source_type == "table":
            cached = session.last_results.get("last_table")
            if cached is None:
                return {
                    "success": False,
                    "error": "No cached table found. Run an analysis first to generate data.",
                }
        elif input_data.source_type == "chart":
            cached = session.last_results.get("last_figure")
            if cached is None:
                return {
                    "success": False,
                    "error": "No cached chart found. Generate a chart first.",
                }
        else:
            return {
                "success": False,
                "error": f"Unknown source_type: {input_data.source_type}",
            }

        # Transform to target format
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = Path(session.config.output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        try:
            if input_data.source_type == "table":
                import pandas as pd
                df = cached
                
                if input_data.target_format == "csv":
                    output_path = output_dir / f"transformed_{timestamp}.csv"
                    df.to_csv(output_path, index=False)
                elif input_data.target_format in ["excel", "xlsx"]:
                    output_path = output_dir / f"transformed_{timestamp}.xlsx"
                    df.to_excel(output_path, index=False, engine="openpyxl")
                elif input_data.target_format == "json":
                    output_path = output_dir / f"transformed_{timestamp}.json"
                    df.to_json(output_path, orient="records", indent=2)
                elif input_data.target_format in ["pdf", "docx"]:
                    # Simple text-based report
                    output_path = output_dir / f"report_{timestamp}.{input_data.target_format}"
                    _create_report(df, output_path, input_data.target_format)
                else:
                    return {
                        "success": False,
                        "error": f"Unsupported format for tables: {input_data.target_format}",
                    }
            
            elif input_data.source_type == "chart":
                fig = cached
                
                if input_data.target_format == "png":
                    output_path = output_dir / f"chart_{timestamp}.png"
                    fig.savefig(output_path, dpi=150, bbox_inches="tight")
                else:
                    return {
                        "success": False,
                        "error": f"Unsupported format for charts: {input_data.target_format}. Only PNG is supported.",
                    }

            return {
                "success": True,
                "description": input_data.description,
                "source_type": input_data.source_type,
                "target_format": input_data.target_format,
                "output_path": str(output_path),
                "message": f"Successfully converted cached {input_data.source_type} to {input_data.target_format}",
            }

        except Exception as e:
            return {
                "success": False,
                "error": f"Transformation failed: {e}",
            }


def _create_report(df, output_path: Path, format_type: str) -> None:
    """Create a simple report from a DataFrame."""
    if format_type == "pdf":
        from fpdf import FPDF
        
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, "Data Analysis Report", ln=True, align="C")
        pdf.ln(10)
        
        pdf.set_font("Arial", "", 10)
        pdf.cell(0, 10, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True)
        pdf.ln(5)
        
        # Add table data
        pdf.set_font("Arial", "B", 10)
        for col in df.columns:
            pdf.cell(40, 10, str(col)[:15], border=1)
        pdf.ln()
        
        pdf.set_font("Arial", "", 9)
        for _, row in df.head(20).iterrows():  # Limit to 20 rows
            for val in row:
                pdf.cell(40, 10, str(val)[:15], border=1)
            pdf.ln()
        
        pdf.output(str(output_path))
    
    elif format_type == "docx":
        from docx import Document
        
        doc = Document()
        doc.add_heading("Data Analysis Report", 0)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Add table
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = "Light Grid Accent 1"
        
        # Header row
        header_cells = table.rows[0].cells
        for idx, col in enumerate(df.columns):
            header_cells[idx].text = str(col)
        
        # Data rows (limit to 50)
        for _, row in df.head(50).iterrows():
            row_cells = table.add_row().cells
            for idx, val in enumerate(row):
                row_cells[idx].text = str(val)
        
        doc.save(str(output_path))

